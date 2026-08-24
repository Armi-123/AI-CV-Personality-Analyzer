# ============================================================
# DOCX CV PARSER
# AI CV Personality Analyzer
# ============================================================

"""
This module extracts text from DOCX CV/resume files.

Responsibilities:
    1. Validate the uploaded DOCX file.
    2. Open the DOCX document safely.
    3. Extract text from paragraphs.
    4. Extract text from tables.
    5. Combine all extracted content.
    6. Clean unnecessary whitespace.
    7. Return the extracted CV text.

This module does NOT:
    - Predict personality traits.
    - Extract skills.
    - Perform career matching.
    - Generate reports.

Those responsibilities are handled by separate modules.
"""


# ============================================================
# 1. IMPORT REQUIRED LIBRARIES
# ============================================================

# Import regular expressions for text cleaning
import re

# Import Document for reading DOCX files
from docx import Document


# ============================================================
# 2. EXTRACT TEXT FROM DOCX
# ============================================================

def extract_docx_text(docx_file):
    """
    Extract text from a DOCX CV/resume.

    Parameters
    ----------
    docx_file : str or file-like object
        Path to the DOCX file or an uploaded DOCX file object.

    Returns
    -------
    str
        Extracted and cleaned CV text.

    Raises
    ------
    ValueError
        If the file is empty, invalid, or contains no readable text.
    """

    # --------------------------------------------------------
    # Validate the input
    # --------------------------------------------------------

    if docx_file is None:

        raise ValueError(
            "No DOCX file was provided."
        )

    try:

        # ----------------------------------------------------
        # Create a DOCX document object
        # ----------------------------------------------------

        document = Document(
            docx_file
        )

    except Exception as error:

        raise ValueError(
            "Unable to read the DOCX file. "
            "Please upload a valid DOCX CV."
        ) from error

    # --------------------------------------------------------
    # Store extracted text
    # --------------------------------------------------------

    extracted_content = []

    # ========================================================
    # 3. EXTRACT TEXT FROM PARAGRAPHS
    # ========================================================

    for paragraph in document.paragraphs:

        # Get paragraph text
        paragraph_text = paragraph.text.strip()

        # Keep only non-empty paragraphs
        if paragraph_text:

            extracted_content.append(
                paragraph_text
            )

    # ========================================================
    # 4. EXTRACT TEXT FROM TABLES
    # ========================================================

    # Some CVs store important information inside tables.
    # Therefore, table content must also be extracted.

    for table in document.tables:

        for row in table.rows:

            row_content = []

            for cell in row.cells:

                # Extract text from the current cell
                cell_text = cell.text.strip()

                # Add non-empty cell content
                if cell_text:

                    row_content.append(
                        cell_text
                    )

            # Combine cells from the current row
            if row_content:

                extracted_content.append(
                    " | ".join(row_content)
                )

    # --------------------------------------------------------
    # Check whether any text was extracted
    # --------------------------------------------------------

    if not extracted_content:

        raise ValueError(
            "No readable text was found in the DOCX file."
        )

    # ========================================================
    # 5. COMBINE EXTRACTED CONTENT
    # ========================================================

    full_text = "\n".join(
        extracted_content
    )

    # ========================================================
    # 6. CLEAN EXTRACTED TEXT
    # ========================================================

    full_text = clean_docx_text(
        full_text
    )

    # --------------------------------------------------------
    # Final validation
    # --------------------------------------------------------

    if not full_text.strip():

        raise ValueError(
            "The DOCX file was read successfully, "
            "but no usable text was found."
        )

    return full_text


# ============================================================
# 7. CLEAN EXTRACTED DOCX TEXT
# ============================================================

def clean_docx_text(text):
    """
    Clean text extracted from a DOCX file.

    Cleaning includes:
        - Removing non-printable characters.
        - Normalizing line breaks.
        - Removing excessive spaces.
        - Removing excessive blank lines.
        - Removing unnecessary leading/trailing spaces.

    Parameters
    ----------
    text : str
        Raw extracted DOCX text.

    Returns
    -------
    str
        Cleaned text.
    """

    # --------------------------------------------------------
    # Validate input
    # --------------------------------------------------------

    if not text:

        return ""

    # Convert input to string
    text = str(text)

    # --------------------------------------------------------
    # Remove non-printable control characters
    # --------------------------------------------------------

    text = re.sub(
        r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]",
        " ",
        text
    )

    # --------------------------------------------------------
    # Normalize line breaks
    # --------------------------------------------------------

    text = text.replace(
        "\r\n",
        "\n"
    )

    text = text.replace(
        "\r",
        "\n"
    )

    # --------------------------------------------------------
    # Remove spaces before line breaks
    # --------------------------------------------------------

    text = re.sub(
        r"[ \t]+\n",
        "\n",
        text
    )

    # --------------------------------------------------------
    # Normalize multiple spaces/tabs
    # --------------------------------------------------------

    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    # --------------------------------------------------------
    # Remove excessive blank lines
    # --------------------------------------------------------

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    # --------------------------------------------------------
    # Remove unnecessary spaces
    # --------------------------------------------------------

    text = text.strip()

    return text


# ============================================================
# 8. VALIDATE DOCX FILE
# ============================================================

def validate_docx_file(docx_file):
    """
    Validate whether the supplied file can be opened as a DOCX.

    Parameters
    ----------
    docx_file : str or file-like object
        DOCX file path or uploaded DOCX object.

    Returns
    -------
    bool
        True when the DOCX file is valid.
    """

    # Check whether a file was provided
    if docx_file is None:

        return False

    try:

        # Try opening the DOCX file
        document = Document(
            docx_file
        )

        # Check whether the document contains
        # paragraphs or tables
        has_paragraphs = len(
            document.paragraphs
        ) > 0

        has_tables = len(
            document.tables
        ) > 0

        return (
            has_paragraphs
            or has_tables
        )

    except Exception:

        return False


# ============================================================
# 9. GET DOCX BASIC INFORMATION
# ============================================================

def get_docx_info(docx_file):
    """
    Get basic information about a DOCX CV.

    Parameters
    ----------
    docx_file : str or file-like object
        DOCX file path or uploaded DOCX object.

    Returns
    -------
    dict
        Basic document information.
    """

    if docx_file is None:

        raise ValueError(
            "No DOCX file was provided."
        )

    try:

        # Open the DOCX document
        document = Document(
            docx_file
        )

        # Count paragraphs
        paragraph_count = len(
            document.paragraphs
        )

        # Count tables
        table_count = len(
            document.tables
        )

        return {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "is_valid": True
        }

    except Exception as error:

        raise ValueError(
            "Unable to read DOCX document information."
        ) from error


# ============================================================
# 10. MODULE TEST
# ============================================================

if __name__ == "__main__":

    print(
        "DOCX parser module loaded successfully."
    )

    print(
        "Use extract_docx_text() to extract text "
        "from a CV DOCX file."
    )